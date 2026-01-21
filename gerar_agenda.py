#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para gerar automaticamente a Agenda da Federação de SAFs
a partir de um arquivo JSON estruturado.
Formato: Folder vertical com 2 colunas e linha divisória no centro.
"""

import json
import os
import re
from datetime import datetime

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Inches, Pt, RGBColor, Twips
except ImportError:
    print("Instalando python-docx...")
    import subprocess

    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Inches, Pt, RGBColor, Twips


def configurar_colunas_secao(
    section, num_colunas=2, espacamento=0.3, linha_divisoria=True
):
    """
    Configura a seção para ter múltiplas colunas com linha divisória preta.
    """
    sectPr = section._sectPr

    # Remover configuração de colunas existente
    cols_existente = sectPr.find(qn('w:cols'))
    if cols_existente is not None:
        sectPr.remove(cols_existente)

    # Criar novo elemento de colunas
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num_colunas))
    cols.set(qn('w:space'), str(int(espacamento * 1440)))

    if linha_divisoria:
        # Ativar linha divisória (w:sep='1' cria uma linha preta entre as colunas)
        cols.set(qn('w:sep'), '1')

    sectPr.append(cols)


def adicionar_linha_vertical_pagina(section):
    """
    Adiciona uma linha vertical preta central no header, cobrindo toda a altura da página.
    Como está no header, aparecerá em todas as páginas automaticamente.
    """

    def _to_pt(valor):
        """Converte para pontos"""
        try:
            return valor.pt
        except Exception:
            # python-docx usa EMU (1 pt = 12700 EMU)
            return valor / 12700

    # Dimensões da página em pontos
    page_height_pt = _to_pt(section.page_height)
    page_width_pt = _to_pt(section.page_width)
    left_margin_pt = _to_pt(section.left_margin)
    right_margin_pt = _to_pt(section.right_margin)

    # Calcular posição X no centro da página
    largura_util = page_width_pt - left_margin_pt - right_margin_pt
    pos_x_centro = left_margin_pt + (largura_util / 2)

    # Largura da linha em pontos (3pt para ficar bem visível e preta)
    largura_linha_pt = 3

    # Criar shape retangular muito fino (linha vertical preta)
    # Usar rect ao invés de line para garantir cor preta sólida
    shape_xml = (
        f'<w:pict {nsdecls("w")} xmlns:v="urn:schemas-microsoft-com:vml">'
        f'<v:rect style="position:absolute;left:{pos_x_centro - largura_linha_pt/2}pt;'
        f'top:0pt;width:{largura_linha_pt}pt;height:{page_height_pt}pt;'
        f'mso-position-horizontal-relative:page;'
        f'mso-position-vertical-relative:page;'
        f'mso-wrap-style:none;z-index:251658240" '
        f'fillcolor="#000000" strokecolor="#000000" strokeweight="0pt"/>'
        f'</w:pict>'
    )

    # Adicionar no header da seção (aparece em todas as páginas)
    header = section.header
    # Limpar parágrafos existentes do header para evitar conflitos
    for para in header.paragraphs:
        para.clear()

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    pict = parse_xml(shape_xml)
    r = p.add_run()
    r._r.append(pict)


def adicionar_titulo_secao(doc, texto):
    """Adiciona um título de seção principal (centralizado, negrito, Arial 14)"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def adicionar_subtitulo(doc, texto):
    """Adiciona um subtítulo (negrito, alinhado à esquerda, Arial 12)"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def adicionar_texto(doc, texto, tamanho=12, negrito=False, italico=False):
    """Adiciona texto compacto (Arial 12)"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(tamanho)
    run.font.name = 'Arial'
    run.bold = negrito
    run.italic = italico
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p


def adicionar_linha(doc, texto, tamanho=12):
    """Adiciona uma linha de texto (Arial 12)"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(tamanho)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


def adicionar_espaco(doc, pts=6):
    """Adiciona um espaço vertical controlado"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    return p


def obter_caminho_imagem(caminho_imagem):
    """Retorna o caminho absoluto da imagem se existir, senão None"""
    if not caminho_imagem:
        return None

    if os.path.isabs(caminho_imagem):
        if os.path.exists(caminho_imagem):
            return caminho_imagem
        return None

    try:
        pasta_atual = os.path.dirname(os.path.abspath(data_file_path))
    except:
        pasta_atual = os.path.dirname(os.path.abspath('agenda_data.json'))

    pasta_fotos = os.path.join(pasta_atual, 'fotos')

    caminhos_tentar = [
        os.path.join(pasta_fotos, caminho_imagem),
        os.path.join(pasta_atual, caminho_imagem),
        caminho_imagem,
    ]

    for caminho in caminhos_tentar:
        if os.path.exists(caminho):
            return caminho

    return None


def adicionar_item_com_foto(doc, foto_path, linhas_texto, largura_foto=0.6):
    """
    Adiciona um item com foto à esquerda e texto à direita.
    linhas_texto: lista de tuplas (texto, tamanho, negrito)
    """
    caminho_foto = obter_caminho_imagem(foto_path) if foto_path else None

    if caminho_foto:
        # Criar tabela de 2 colunas (foto | texto)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Remover bordas da tabela
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

        # Configurar larguras das colunas
        cell_foto = table.rows[0].cells[0]
        cell_texto = table.rows[0].cells[1]

        cell_foto.width = Inches(largura_foto + 0.1)
        cell_texto.width = Inches(4.0)

        # Adicionar foto na célula esquerda
        p_foto = cell_foto.paragraphs[0]
        p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_foto.add_run()

        largura = Inches(largura_foto)
        altura = Inches(largura_foto * (4 / 3))  # Proporção 3x4
        run.add_picture(caminho_foto, width=largura, height=altura)

        # Adicionar texto na célula direita
        first_line = True
        for linha_info in linhas_texto:
            if isinstance(linha_info, tuple):
                texto, tamanho, negrito = linha_info
            else:
                texto, tamanho, negrito = linha_info, 12, False

            if first_line:
                p = cell_texto.paragraphs[0]
                first_line = False
            else:
                p = cell_texto.add_paragraph()

            run = p.add_run(texto)
            run.font.size = Pt(tamanho)
            run.font.name = 'Arial'
            run.bold = negrito
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)

        return table
    else:
        # Sem foto - apenas texto normal
        for linha_info in linhas_texto:
            if isinstance(linha_info, tuple):
                texto, tamanho, negrito = linha_info
            else:
                texto, tamanho, negrito = linha_info, 12, False

            p = doc.add_paragraph()
            run = p.add_run(texto)
            run.font.size = Pt(tamanho)
            run.font.name = 'Arial'
            run.bold = negrito
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)

        return None


def adicionar_imagem(doc, caminho_imagem, largura_base=None):
    """
    Adiciona uma imagem ao documento Word no formato 3x4.
    """
    if not caminho_imagem:
        return False

    try:
        caminho_final = None

        if os.path.isabs(caminho_imagem):
            caminho_final = caminho_imagem
        else:
            try:
                pasta_atual = os.path.dirname(os.path.abspath(data_file_path))
            except:
                pasta_atual = os.path.dirname(os.path.abspath('agenda_data.json'))

            pasta_fotos = os.path.join(pasta_atual, 'fotos')

            caminhos_tentar = [
                os.path.join(pasta_fotos, caminho_imagem),
                os.path.join(pasta_atual, caminho_imagem),
                caminho_imagem,
            ]

            for caminho in caminhos_tentar:
                if os.path.exists(caminho):
                    caminho_final = caminho
                    break

        if not caminho_final or not os.path.exists(caminho_final):
            return False

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        if largura_base is None:
            largura_base = Inches(0.7)

        # Proporção 3x4 (largura:altura = 3:4)
        largura = largura_base
        altura = Inches(largura_base.inches * (4 / 3))

        run = p.add_run()
        run.add_picture(caminho_final, width=largura, height=altura)
        return True
    except Exception as e:
        return False


def adicionar_capa(doc, caminho_capa):
    """
    Adiciona a capa na primeira página (meia folha, uma coluna, sem linha vertical).
    """
    if not caminho_capa:
        return False

    caminho_final = obter_caminho_imagem(caminho_capa)
    if not caminho_final or not os.path.exists(caminho_final):
        return False

    # Configurar primeira seção para 1 coluna (meia folha) sem linha divisória
    section = doc.sections[0]
    configurar_colunas_secao(
        section, num_colunas=1, espacamento=0.25, linha_divisoria=False
    )
    # NÃO adicionar linha vertical no header (essa página não deve ter)

    # Adicionar imagem da capa alinhada à esquerda (mesmo tamanho do calendário)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Calcular tamanho IDÊNTICO ao calendário
    # Mesmas dimensões usadas no calendário na última página
    largura_total = Inches(11.69 - 0.8)  # Largura útil da página
    largura_capa = Inches((largura_total.inches / 2) - 0.3)  # Igual ao calendário

    run = p.add_run()
    try:
        # Usar exatamente a mesma largura do calendário (mantém proporção)
        run.add_picture(caminho_final, width=largura_capa)
    except:
        # Se falhar, tentar novamente
        run.add_picture(caminho_final, width=largura_capa)

    return True


def adicionar_ultima_pagina(doc, caminho_calendario):
    """
    Adiciona a última página com calendário à esquerda e anotações gerais à direita.
    Sem linha vertical.
    """
    # Criar nova seção para a última página
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Configurar margens
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    # NÃO configurar colunas (vamos usar tabela para layout lado a lado)
    # NÃO adicionar linha vertical no header (essa página não deve ter)

    # Criar tabela de 2 colunas: calendário à esquerda, anotações à direita
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Remover bordas da tabela
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # Configurar larguras das colunas (50% cada)
    cell_calendario = table.rows[0].cells[0]
    cell_anotacoes = table.rows[0].cells[1]

    largura_total = Inches(11.69 - 0.8)  # Largura útil da página
    cell_calendario.width = Inches(largura_total.inches / 2)
    cell_anotacoes.width = Inches(largura_total.inches / 2)

    # === COLUNA ESQUERDA: CALENDÁRIO ===
    if caminho_calendario:
        caminho_final = obter_caminho_imagem(caminho_calendario)
        if caminho_final and os.path.exists(caminho_final):
            p_calendario = cell_calendario.paragraphs[0]
            p_calendario.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_calendario.paragraph_format.space_before = Pt(0)
            p_calendario.paragraph_format.space_after = Pt(0)

            # Calcular tamanho para a célula esquerda (meia página)
            # Largura da célula menos margem
            largura_calendario = Inches((largura_total.inches / 2) - 0.3)
            # Altura proporcional (não ocupar toda a página para não empurrar anotações)
            altura_calendario = Inches(7.0)

            run = p_calendario.add_run()
            try:
                # Usar largura para manter proporção correta
                run.add_picture(caminho_final, width=largura_calendario)
            except:
                run.add_picture(caminho_final, width=largura_calendario)
    else:
        # Se não houver calendário, deixar célula esquerda vazia
        p_vazio = cell_calendario.paragraphs[0]
        p_vazio.clear()

    # === COLUNA DIREITA: ANOTAÇÕES GERAIS ===
    # Sempre adicionar anotações gerais, mesmo sem calendário
    # Garantir que a célula tenha pelo menos um parágrafo
    if len(cell_anotacoes.paragraphs) == 0:
        cell_anotacoes.add_paragraph()

    # Limpar parágrafo padrão da célula
    p_titulo = cell_anotacoes.paragraphs[0]
    p_titulo.clear()

    # Adicionar título "Anotações gerais" (centralizado)
    run_titulo = p_titulo.add_run("Anotações gerais")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    run_titulo.font.name = 'Arial'
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(0)
    p_titulo.paragraph_format.space_after = Pt(15)

    # Calcular quantas linhas cabem na célula
    altura_util_pt = (8.27 - 0.8) * 72  # Altura útil em pontos (7.47" = 537.84pt)
    altura_usada_pt = 14 + 15  # Título (14pt) e espaçamento (15pt)
    altura_disponivel_pt = altura_util_pt - altura_usada_pt

    # Altura de cada linha (aproximadamente 20pt com espaçamento)
    altura_linha_pt = 20
    num_linhas = max(
        15, int(altura_disponivel_pt / altura_linha_pt)
    )  # Mínimo 15 linhas

    # Garantir que sempre adicionamos pelo menos algumas linhas
    if num_linhas < 15:
        num_linhas = 25  # Mínimo de 25 linhas se o cálculo der errado

    # Adicionar linhas horizontais centralizadas até o fim da célula
    for i in range(num_linhas):
        try:
            p = cell_anotacoes.add_paragraph()
            # Linha horizontal simples centralizada
            run = p.add_run("_" * 50)  # Linha de sublinhado ajustada para meia folha
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centralizar linhas
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.0
        except Exception as e:
            # Se houver erro, continuar tentando
            print(f"Erro ao adicionar linha {i}: {e}")
            continue


# Variável global para armazenar o caminho do arquivo de dados
data_file_path = 'agenda_data.json'


def gerar_agenda(data_file='agenda_data.json', output_file=None):
    """Gera o documento Word da agenda a partir do JSON"""

    global data_file_path
    data_file_path = os.path.abspath(data_file)

    # Carregar dados
    with open(data_file, 'r', encoding='utf-8') as f:
        dados = json.load(f)

    doc = Document()

    # Configurar página - A4 paisagem
    section = doc.sections[0]
    section.page_height = Inches(8.27)
    section.page_width = Inches(11.69)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    # ==================== CAPA ====================
    # Adicionar capa na primeira página (se existir)
    if dados.get('capa'):
        adicionar_capa(doc, dados.get('capa'))
        # Após a capa, criar nova seção para o conteúdo com colunas
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.page_height = Inches(8.27)
        section.page_width = Inches(11.69)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
    else:
        # Se não houver capa, usar a primeira seção normalmente
        pass

    # CONFIGURAR 2 COLUNAS COM LINHA DIVISÓRIA NO CENTRO
    configurar_colunas_secao(
        section, num_colunas=2, espacamento=0.25, linha_divisoria=True
    )
    # Adicionar linha vertical preta no header (aparece em todas as páginas)
    adicionar_linha_vertical_pagina(section)

    # ==================== CONTEÚDO ====================

    # === PALAVRA DA PRESIDENTE ===
    adicionar_titulo_secao(doc, "PALAVRA DA PRESIDENTE")

    mensagem = dados['presidente']['mensagem'].replace('\n\n', '\n').strip()
    adicionar_texto(doc, mensagem, tamanho=10)
    adicionar_texto(doc, dados['presidente']['nome'], tamanho=10, negrito=True)
    adicionar_espaco(doc, 8)

    # === DIRETORIA ===
    adicionar_titulo_secao(doc, "I - DIRETORIA")

    for membro in dados['diretoria']:
        if membro.get('nome'):
            # Montar linhas de texto
            linhas = []

            # Cargo e Nome
            cargo_nome = f"{membro['cargo']}: {membro['nome']}"
            if membro.get('data_nascimento'):
                cargo_nome += f" (DN: {membro['data_nascimento']})"
            linhas.append((cargo_nome, 12, True))

            # Email e endereço
            if membro.get('email'):
                linhas.append((f"Email: {membro['email']}", 10, False))
            if membro.get('endereco'):
                linhas.append((f"End: {membro['endereco']}", 10, False))

            # Adicionar com foto à esquerda
            adicionar_item_com_foto(doc, membro.get('foto'), linhas, largura_foto=0.5)
            adicionar_espaco(doc, 2)

    adicionar_espaco(doc, 6)

    # === SAFs ===
    adicionar_titulo_secao(doc, "II - SAFs FILIADAS")

    for saf in dados['safs']:
        # Título da SAF
        adicionar_subtitulo(doc, f"{saf['numero']}. {saf['nome']}")

        # Montar linhas de texto
        linhas = []

        # Endereço
        linhas.append((f"End: {saf['endereco']}", 11, False))

        # Pastor
        if saf.get('pastor') and saf['pastor'].get('nome'):
            pastor = saf['pastor']
            linha = f"Pastor: {pastor['nome']}"
            if pastor.get('data_nascimento'):
                linha += f" ({pastor['data_nascimento']})"
            linhas.append((linha, 11, False))

        # Presidente
        if saf.get('presidente') and saf['presidente'].get('nome'):
            pres = saf['presidente']
            linha_pres = f"Presidente: {pres['nome']}"
            if pres.get('data_nascimento'):
                linha_pres += f" ({pres['data_nascimento']})"
            linhas.append((linha_pres, 11, False))

            # Contato
            contato = []
            if pres.get('telefone'):
                contato.append(pres['telefone'])
            if pres.get('email'):
                contato.append(pres['email'])
            if contato:
                linhas.append(("Contato: " + " | ".join(contato), 10, False))

        # Conselheiro
        if saf.get('conselheiro') and saf['conselheiro'].get('nome'):
            cons = saf['conselheiro']
            linha_cons = f"Conselheiro: {cons['nome']}"
            if cons.get('data_nascimento'):
                linha_cons += f" ({cons['data_nascimento']})"
            linhas.append((linha_cons, 11, False))

        # Aniversário
        if saf.get('aniversario'):
            aniv = saf['aniversario']
            linhas.append(
                (f"Aniversário: {aniv['data']} - {aniv['anos']} anos", 11, False)
            )

        # Adicionar com foto à esquerda
        adicionar_item_com_foto(doc, saf.get('foto'), linhas, largura_foto=0.6)
        adicionar_espaco(doc, 4)

    # === ATIVIDADES REALIZADAS ===
    adicionar_titulo_secao(doc, "III - ATIVIDADES REALIZADAS EM 2023")

    for atividade in dados.get('atividades_realizadas_2023', []):
        linha = atividade['descricao']
        if atividade.get('data'):
            linha = f"{atividade['data']} - {linha}"
        adicionar_linha(doc, f"• {linha}")

    adicionar_espaco(doc, 8)

    # === ATIVIDADES PLANEJADAS ===
    ano_atual = dados.get('ano', 2024)
    adicionar_titulo_secao(doc, f"IV - ATIVIDADES PLANEJADAS PARA {ano_atual}")

    meses_pt = {
        'janeiro': 'JANEIRO',
        'fevereiro': 'FEVEREIRO',
        'marco': 'MARÇO',
        'abril': 'ABRIL',
        'maio': 'MAIO',
        'junho': 'JUNHO',
        'julho': 'JULHO',
        'agosto': 'AGOSTO',
        'setembro': 'SETEMBRO',
        'outubro': 'OUTUBRO',
        'novembro': 'NOVEMBRO',
        'dezembro': 'DEZEMBRO',
    }

    chave_atividades = f'atividades_planejadas_{ano_atual}'
    atividades_planejadas = dados.get(chave_atividades, {})

    for mes_key, mes_nome in meses_pt.items():
        if mes_key in atividades_planejadas and atividades_planejadas[mes_key]:
            adicionar_subtitulo(doc, mes_nome)

            for atividade in atividades_planejadas[mes_key]:
                linha = atividade['descricao']
                if atividade.get('data'):
                    linha = f"{atividade['data']} - {linha}"
                adicionar_linha(doc, f"• {linha}")

    adicionar_espaco(doc, 8)

    # === INFORMAÇÕES GERAIS ===
    adicionar_titulo_secao(doc, "V - INFORMAÇÕES GERAIS")

    if dados.get('informacoes_gerais'):
        info = dados['informacoes_gerais']

        # Missionário de Oração
        if info.get('missionario_oracao'):
            miss = info['missionario_oracao']
            adicionar_subtitulo(doc, "Missionário de Oração:")

            linhas_miss = [
                (
                    f"Rev. {miss.get('nome', '')} (DN: {miss.get('data_nascimento', '')})",
                    12,
                    True,
                ),
                (f"Campo: {miss.get('campo', '')}", 11, False),
                (f"WhatsApp: {miss.get('whatsapp', '')}", 11, False),
            ]

            adicionar_item_com_foto(
                doc, miss.get('foto'), linhas_miss, largura_foto=0.6
            )
            adicionar_espaco(doc, 4)

        # Observações
        if info.get('observacoes'):
            adicionar_subtitulo(doc, "Observações:")
            for obs in info['observacoes']:
                adicionar_linha(doc, f"• {obs}")
            adicionar_espaco(doc, 4)

        # Lema
        if info.get('lema'):
            adicionar_subtitulo(doc, "Lema:")
            for linha in info['lema']:
                p = adicionar_linha(doc, linha)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ==================== ÚLTIMA PÁGINA ====================
    # Adicionar última página com calendário e anotações gerais
    # Sempre adicionar, mesmo sem calendário
    calendario_path = dados.get('calendario', '')
    print(f"[DEBUG] Adicionando última página. Calendário: {calendario_path}")
    adicionar_ultima_pagina(doc, calendario_path)
    print(f"[DEBUG] Última página adicionada com sucesso")

    # Salvar documento - sempre usar o ano do JSON
    ano = dados.get('ano', 2024)

    if output_file is None:
        output_file = f"Agenda {ano}.docx"
    else:
        # Substituir qualquer ano no nome do arquivo pelo ano do JSON
        nome_base, ext = os.path.splitext(output_file)
        # Remover qualquer ano existente (4 dígitos)
        nome_base = re.sub(r'\s*\d{4}\s*', f' {ano} ', nome_base)
        nome_base = nome_base.strip()
        # Garantir que o ano está no nome
        if str(ano) not in nome_base:
            nome_base = f"{nome_base} {ano}"
        output_file = f"{nome_base}{ext}"

    doc.save(output_file)
    print(f"[OK] Documento gerado com sucesso: {output_file}")


if __name__ == "__main__":
    import sys

    data_file = sys.argv[1] if len(sys.argv) > 1 else 'agenda_data.json'
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        gerar_agenda(data_file, output_file)
    except FileNotFoundError:
        print(f"Erro: Arquivo '{data_file}' não encontrado!")
    except Exception as e:
        print(f"Erro ao gerar agenda: {e}")
        import traceback

        traceback.print_exc()
