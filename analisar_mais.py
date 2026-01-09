#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Analisa mais detalhes do documento Word"""

from docx import Document

doc = Document("Agenda número 2023 (1).docx")

# Procurar por "SAF" ou "Saf"
print("PROCURANDO SECOES SAF:")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    texto = para.text.strip()
    if 'SAF' in texto.upper() or 'Saf' in texto or 'saf' in texto.lower():
        print(f"\nParagrafo {i+1}: {texto[:150]}")
        if para.runs:
            print(
                f"  Fonte: {para.runs[0].font.name if para.runs[0].font.name else 'N/A'}"
            )
            print(
                f"  Tamanho: {para.runs[0].font.size / 12700 if para.runs[0].font.size else 'N/A'} pt"
            )
            print(f"  Negrito: {para.runs[0].font.bold}")

# Procurar por "Informacoes Gerais" ou "Missionario"
print("\n\nPROCURANDO SECOES DE INFORMACOES:")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    texto = para.text.strip()
    if any(
        palavra in texto.lower()
        for palavra in ['informacao', 'missionario', 'observacao', 'lema']
    ):
        print(f"\nParagrafo {i+1}: {texto[:150]}")
        if para.runs:
            print(
                f"  Fonte: {para.runs[0].font.name if para.runs[0].font.name else 'N/A'}"
            )
            print(
                f"  Tamanho: {para.runs[0].font.size / 12700 if para.runs[0].font.size else 'N/A'} pt"
            )

# Contar parágrafos totais
print(f"\n\nTOTAL DE PARAGRAFOS: {len(doc.paragraphs)}")
