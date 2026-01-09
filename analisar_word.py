#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para analisar um documento Word existente e extrair seu formato
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def analisar_documento(arquivo):
    """Analisa um documento Word e extrai informações de formatação"""
    doc = Document(arquivo)

    print("=" * 80)
    print(f"ANÁLISE DO DOCUMENTO: {arquivo}")
    print("=" * 80)

    # Informações da seção
    section = doc.sections[0]
    print(f"\nCONFIGURACAO DA PAGINA:")
    print(
        f"  Altura: {section.page_height / 914400:.2f}\" ({section.page_height / 914400 * 2.54:.2f} cm)"
    )
    print(
        f"  Largura: {section.page_width / 914400:.2f}\" ({section.page_width / 914400 * 2.54:.2f} cm)"
    )
    print(f"  Margem esquerda: {section.left_margin / 914400:.2f}\"")
    print(f"  Margem direita: {section.right_margin / 914400:.2f}\"")
    print(f"  Margem superior: {section.top_margin / 914400:.2f}\"")
    print(f"  Margem inferior: {section.bottom_margin / 914400:.2f}\"")

    # Analisar parágrafos
    print(f"\nPARAGRAFOS ({len(doc.paragraphs)}):")
    print("-" * 80)

    for i, para in enumerate(doc.paragraphs[:30]):  # Primeiros 30 parágrafos
        texto = para.text.strip()
        if not texto:
            continue

        # Informações do parágrafo
        alinhamento = para.alignment
        alinhamento_str = (
            {0: "ESQUERDA", 1: "CENTRO", 2: "DIREITA", 3: "JUSTIFICADO"}.get(
                alinhamento, "DESCONHECIDO"
            )
            if alinhamento is not None
            else "None"
        )

        # Informações do formato
        if para.runs:
            first_run = para.runs[0]
            tamanho = first_run.font.size / 12700 if first_run.font.size else "N/A"
            negrito = first_run.font.bold
            italico = first_run.font.italic
            fonte = first_run.font.name if first_run.font.name else "N/A"
        else:
            tamanho = "N/A"
            negrito = False
            italico = False
            fonte = "N/A"

        # Espaçamento
        space_before = (
            para.paragraph_format.space_before / 12700
            if para.paragraph_format.space_before
            else 0
        )
        space_after = (
            para.paragraph_format.space_after / 12700
            if para.paragraph_format.space_after
            else 0
        )

        print(f"\nParágrafo {i+1}:")
        print(f"  Texto: {texto[:100]}{'...' if len(texto) > 100 else ''}")
        print(f"  Alinhamento: {alinhamento_str}")
        print(f"  Tamanho da fonte: {tamanho} pt")
        print(f"  Fonte: {fonte}")
        print(f"  Negrito: {negrito}")
        print(f"  Itálico: {italico}")
        print(f"  Espaço antes: {space_before:.1f} pt")
        print(f"  Espaço depois: {space_after:.1f} pt")

        # Verificar imagens
        try:
            if para._element.xpath('.//a:blip'):
                print(f"  [CONTEM IMAGEM]")
        except:
            pass

    # Verificar tabelas
    print(f"\nTABELAS ({len(doc.tables)}):")
    for i, table in enumerate(doc.tables):
        print(
            f"  Tabela {i+1}: {len(table.rows)} linhas x {len(table.columns)} colunas"
        )

    # Estilos
    print(f"\nESTILOS DISPONIVEIS:")
    for style in doc.styles[:10]:
        if style.type == 1:  # Parágrafo
            print(f"  - {style.name}")


if __name__ == "__main__":
    arquivo = "Agenda número 2023 (1).docx"
    try:
        analisar_documento(arquivo)
    except Exception as e:
        print(f"Erro: {e}")
        import traceback

        traceback.print_exc()
