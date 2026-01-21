#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de teste para verificar se as anotações aparecem corretamente
"""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

doc = Document()

# Configurar página - A4 paisagem
section = doc.sections[0]
section.page_height = Inches(8.27)
section.page_width = Inches(11.69)
section.left_margin = Inches(0.4)
section.right_margin = Inches(0.4)
section.top_margin = Inches(0.4)
section.bottom_margin = Inches(0.4)

# Adicionar conteúdo de teste na primeira página
doc.add_paragraph('Página de teste - conteúdo principal')

# Criar nova seção para a última página
section2 = doc.add_section(WD_SECTION.NEW_PAGE)
section2.left_margin = Inches(0.4)
section2.right_margin = Inches(0.4)
section2.top_margin = Inches(0.4)
section2.bottom_margin = Inches(0.4)

# Criar tabela de 2 colunas
table = doc.add_table(rows=1, cols=2)
table.autofit = False

# Remover bordas da tabela
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
tblBorders = OxmlElement('w:tblBorders')
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'nil')
    tblBorders.append(border)
tblPr.append(tblBorders)
if tbl.tblPr is None:
    tbl.insert(0, tblPr)

cell_esquerda = table.rows[0].cells[0]
cell_direita = table.rows[0].cells[1]

largura_total = Inches(11.69 - 0.8)
cell_esquerda.width = Inches(largura_total.inches / 2)
cell_direita.width = Inches(largura_total.inches / 2)

# Célula esquerda - texto de placeholder (simulando calendário)
p_esq = cell_esquerda.paragraphs[0]
p_esq.add_run('[CALENDÁRIO AQUI - COLUNA ESQUERDA]')

# Célula direita - Anotações gerais
print("Adicionando anotações gerais na célula direita...")

# Limpar e adicionar título
p_titulo = cell_direita.paragraphs[0]
p_titulo.clear()
run_titulo = p_titulo.add_run('Anotações gerais')
run_titulo.bold = True
run_titulo.font.size = Pt(14)
run_titulo.font.name = 'Arial'
p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_titulo.paragraph_format.space_before = Pt(0)
p_titulo.paragraph_format.space_after = Pt(15)

# Adicionar 25 linhas
for i in range(25):
    p = cell_direita.add_paragraph()
    run = p.add_run('_' * 50)
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0

print(f"Total de parágrafos na célula direita: {len(cell_direita.paragraphs)}")

doc.save('teste_anotacoes.docx')
print('Arquivo teste_anotacoes.docx criado com sucesso!')
print('Abra o arquivo para verificar se as anotações aparecem na coluna direita.')
